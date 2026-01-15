"""
Script para transcribir audio en chunks con logging detallado.
Usa prioridad baja para no congelar la PC.
"""
import os
import sys
import time
import json
from pathlib import Path
from datetime import datetime

# Agregar FFmpeg al PATH
os.environ["PATH"] = "C:\\ffmpeg;" + os.environ.get("PATH", "")

# Establecer prioridad baja ANTES de importar torch
if sys.platform == 'win32':
    import ctypes
    BELOW_NORMAL_PRIORITY_CLASS = 0x00004000
    handle = ctypes.windll.kernel32.GetCurrentProcess()
    ctypes.windll.kernel32.SetPriorityClass(handle, BELOW_NORMAL_PRIORITY_CLASS)
    print("[INFO] Prioridad del proceso establecida a BELOW_NORMAL")

# Configurar logging
LOG_FILE = Path("C:/Users/ghell/bout/logs/transcription_chunks.log")
LOG_FILE.parent.mkdir(parents=True, exist_ok=True)

def log(message: str):
    """Log a message to file and console."""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{timestamp}] {message}"
    print(line)
    with open(LOG_FILE, "a", encoding="utf-8") as f:
        f.write(line + "\n")

def main():
    chunks_dir = Path("C:/Users/ghell/bout/temp/chunks")
    output_dir = Path("C:/Users/ghell/bout/output")
    results_file = Path("C:/Users/ghell/bout/temp/transcription_results.json")

    # Listar chunks
    chunks = sorted(chunks_dir.glob("chunk_*.wav"))
    log(f"Encontrados {len(chunks)} chunks para transcribir")

    # Cargar resultados previos si existen (para continuar si se interrumpe)
    results = {}
    if results_file.exists():
        try:
            with open(results_file, "r", encoding="utf-8") as f:
                results = json.load(f)
            log(f"Cargados {len(results)} resultados previos")
        except:
            pass

    # Importar whisper (esto toma tiempo)
    log("Importando whisper...")
    import whisper

    # Cargar modelo
    log("Cargando modelo whisper 'small' en CUDA (menos VRAM)...")
    start_load = time.time()
    model = whisper.load_model("small", device="cuda")
    log(f"Modelo cargado en {time.time() - start_load:.1f} segundos")

    # Transcribir cada chunk
    for i, chunk_path in enumerate(chunks):
        chunk_name = chunk_path.name

        # Saltar si ya fue transcrito
        if chunk_name in results:
            log(f"[{i+1}/{len(chunks)}] {chunk_name} - YA TRANSCRITO, saltando")
            continue

        log(f"[{i+1}/{len(chunks)}] Transcribiendo {chunk_name}...")
        start_time = time.time()

        try:
            result = model.transcribe(
                str(chunk_path),
                language="es",
                task="transcribe",
                verbose=False
            )

            elapsed = time.time() - start_time
            text = result["text"].strip()

            log(f"[{i+1}/{len(chunks)}] {chunk_name} completado en {elapsed:.1f}s - {len(text)} caracteres")

            # Guardar resultado
            results[chunk_name] = {
                "text": text,
                "duration": elapsed,
                "timestamp": datetime.now().isoformat()
            }

            # Guardar resultados intermedios (por si se interrumpe)
            with open(results_file, "w", encoding="utf-8") as f:
                json.dump(results, f, ensure_ascii=False, indent=2)

            # Pausa breve para que el sistema respire
            time.sleep(1)

        except Exception as e:
            log(f"[{i+1}/{len(chunks)}] ERROR en {chunk_name}: {e}")
            continue

    # Combinar todos los resultados
    log("Combinando resultados...")
    full_text = ""
    for chunk_path in chunks:
        chunk_name = chunk_path.name
        if chunk_name in results:
            full_text += results[chunk_name]["text"] + " "

    full_text = full_text.strip()
    log(f"Texto total: {len(full_text)} caracteres")

    # Guardar texto plano
    text_output = output_dir / "(8428-2024) AUDIENCIA_transcripcion_chunks.txt"
    with open(text_output, "w", encoding="utf-8") as f:
        f.write(full_text)
    log(f"Texto guardado en: {text_output}")

    # Crear documento DOCX
    try:
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        # Titulo
        title = doc.add_heading("TRANSCRIPCION DE AUDIENCIA", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Info
        doc.add_paragraph(f"Archivo: (8428-2024) AUDIENCIA CON MENORES DE EDAD 3_12_25.mp4")
        doc.add_paragraph(f"Fecha de transcripcion: {datetime.now().strftime('%d/%m/%Y %H:%M')}")
        doc.add_paragraph(f"Duracion: ~40 minutos")
        doc.add_paragraph("")

        # Contenido
        doc.add_heading("CONTENIDO", level=1)

        # Agregar texto en parrafos
        paragraphs = full_text.split(". ")
        current_para = ""
        for sentence in paragraphs:
            current_para += sentence + ". "
            if len(current_para) > 500:
                p = doc.add_paragraph(current_para.strip())
                p.style.font.size = Pt(11)
                current_para = ""

        if current_para:
            p = doc.add_paragraph(current_para.strip())
            p.style.font.size = Pt(11)

        docx_output = output_dir / "(8428-2024) AUDIENCIA_transcripcion_chunks.docx"
        doc.save(str(docx_output))
        log(f"Documento DOCX guardado en: {docx_output}")

    except Exception as e:
        log(f"Error creando DOCX: {e}")

    log("=== TRANSCRIPCION COMPLETADA ===")
    return 0

if __name__ == "__main__":
    try:
        sys.exit(main())
    except KeyboardInterrupt:
        log("Interrumpido por usuario")
        sys.exit(1)
    except Exception as e:
        log(f"Error fatal: {e}")
        import traceback
        traceback.print_exc()
        sys.exit(1)

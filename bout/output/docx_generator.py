"""
Word document generation for transcriptions.

Creates formatted .docx files with metadata and transcription content.
"""
from datetime import datetime
from pathlib import Path
from typing import List, Optional

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from ..core.types import TranscriptionSegment
from ..core.config import get_config
from ..utils.paths import PathManager
from ..logging import get_logger

logger = get_logger("output.docx")


class DocumentGenerator:
    """
    Generates Word documents from transcriptions.

    Features:
    - Formatted header with metadata
    - Timestamped transcription content
    - Optional speaker labels
    - Clean, professional formatting
    """

    def __init__(self):
        """Initialize document generator."""
        self.config = get_config()

    def generate(
        self,
        video_name: str,
        text: str,
        segments: List[TranscriptionSegment],
        duration_seconds: float,
        output_path: Optional[Path] = None,
    ) -> Path:
        """
        Generate transcription document.

        Args:
            video_name: Original video filename
            text: Full transcription text
            segments: List of transcription segments with timing
            duration_seconds: Total video duration
            output_path: Output file path (auto-generated if None)

        Returns:
            Path to generated document
        """
        # Generate output path
        if output_path is None:
            safe_name = PathManager.safe_filename(Path(video_name).stem)
            output_name = f"{safe_name}_transcripcion.docx"
            output_path = self.config.output_dir / output_name

        output_path = PathManager.normalize(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)

        # Handle existing file
        if output_path.exists():
            output_path = PathManager.get_unique_path(output_path)

        logger.info(f"Generating document: {output_path.name}")

        # Create document
        doc = Document()
        self._setup_styles(doc)

        # Add content
        self._add_header(doc, video_name, duration_seconds)
        self._add_separator(doc)

        if segments:
            self._add_segments(doc, segments)
        else:
            self._add_plain_text(doc, text)

        self._add_footer(doc)

        # Save document
        doc.save(str(output_path))

        logger.info(f"Document saved: {output_path}")
        return output_path

    def _setup_styles(self, doc: Document):
        """Configure document styles."""
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Calibri"
        font.size = Pt(11)

    def _add_header(self, doc: Document, video_name: str, duration_seconds: float):
        """Add document header with metadata."""
        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("TRANSCRIPCION")
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()

        # Metadata table
        table = doc.add_table(rows=4, cols=2)
        table.autofit = True

        rows = [
            ("Archivo original:", video_name),
            ("Fecha de transcripcion:", datetime.now().strftime("%d/%m/%Y %H:%M")),
            ("Duracion del video:", self._format_duration(duration_seconds)),
            ("Procesado con:", "BOUT v2.0"),
        ]

        for i, (label, value) in enumerate(rows):
            cells = table.rows[i].cells
            cells[0].text = label
            cells[1].text = value

            # Bold labels
            for paragraph in cells[0].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

        doc.add_paragraph()

    def _add_separator(self, doc: Document):
        """Add visual separator line."""
        sep = doc.add_paragraph()
        sep.add_run("─" * 70)
        doc.add_paragraph()

    def _add_segments(self, doc: Document, segments: List[TranscriptionSegment]):
        """Add transcription with timestamps."""
        for segment in segments:
            para = doc.add_paragraph()

            # Timestamp
            timestamp = self._format_timestamp(segment.start)
            ts_run = para.add_run(f"[{timestamp}] ")
            ts_run.font.color.rgb = RGBColor(128, 128, 128)
            ts_run.font.size = Pt(9)

            # Speaker (if available)
            if segment.speaker:
                speaker_run = para.add_run(f"{segment.speaker}: ")
                speaker_run.bold = True

            # Text
            para.add_run(segment.text)

            # Add spacing
            para.paragraph_format.space_after = Pt(6)

    def _add_plain_text(self, doc: Document, text: str):
        """Add plain text without timestamps."""
        # Split into paragraphs for readability
        paragraphs = text.split("\n")

        for para_text in paragraphs:
            para_text = para_text.strip()
            if para_text:
                para = doc.add_paragraph(para_text)
                para.paragraph_format.space_after = Pt(6)

    def _add_footer(self, doc: Document):
        """Add document footer."""
        doc.add_paragraph()
        self._add_separator(doc)

        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run(
            f"Documento generado el {datetime.now().strftime('%d/%m/%Y a las %H:%M')}"
        )
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)

    @staticmethod
    def _format_duration(seconds: float) -> str:
        """Format seconds as human-readable duration."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)

        if hours > 0:
            return f"{hours}h {minutes}m {secs}s"
        return f"{minutes}m {secs}s"

    @staticmethod
    def _format_timestamp(seconds: float) -> str:
        """Format seconds as HH:MM:SS timestamp."""
        hours = int(seconds // 3600)
        minutes = int((seconds % 3600) // 60)
        secs = int(seconds % 60)

        if hours > 0:
            return f"{hours:02d}:{minutes:02d}:{secs:02d}"
        return f"{minutes:02d}:{secs:02d}"
